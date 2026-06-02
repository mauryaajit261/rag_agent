"""
mySetu AI - Document Ingestion with Pinecone
Precision chunking, embedding, and vector storage for strict document-grounded RAG.
"""

import os
import re
import uuid
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

# Document loaders
import fitz  # PyMuPDF
from docx import Document
import pandas as pd

from langchain_core.documents import Document as LangChainDocument

from chunker import SmartSemanticChunker

# Pinecone
from pinecone import Pinecone, ServerlessSpec

from config import settings
from models import DocumentType, ProcessingStatus, DocumentInfo

# Set Pinecone API key as environment variable (required by langchain-pinecone)
os.environ['PINECONE_API_KEY'] = settings.PINECONE_API_KEY

# Now import langchain-pinecone (PineconeEmbeddings uses Pinecone's hosted inference)
from langchain_pinecone import PineconeVectorStore, PineconeEmbeddings


class DocumentProcessor:
    """Handles document processing and Pinecone vector storage — precision pipeline"""

    def split_by_headings(self, text: str) -> List[Dict]:
        """
        Split text into sections based on headings.
        Supports:
        - Markdown (## Heading)
        - ALL CAPS headings
        """
        sections = []
        pattern = r'(?:^|\n)(##\s.*|\n[A-Z][A-Z\s]{5,}\n)'
        splits = re.split(pattern, text)
        current_section = "general"

        for part in splits:
            if not part.strip():
                continue
            if part.strip().startswith("##") or part.strip().isupper():
                current_section = part.strip().replace("##", "").strip()
            else:
                sections.append({
                    "section_title": current_section,
                    "content": part.strip()
                })
        return sections

    def __init__(self):
        # Embeddings — Pinecone-hosted inference (no local model, no torch → small footprint).
        # The wrapper sets the e5 query/passage input types automatically.
        self.embeddings = PineconeEmbeddings(
            model=settings.EMBEDDING_MODEL,
            pinecone_api_key=settings.PINECONE_API_KEY,
        )

        # Document metadata
        self.document_metadata: Dict[str, DocumentInfo] = {}

        # Initialize Pinecone
        self._init_pinecone()

        # Load existing metadata
        self._load_metadata()

    def _init_pinecone(self):
        """Initialize Pinecone client and index"""
        try:
            self.pc = Pinecone(api_key=settings.PINECONE_API_KEY)

            # Check if index exists
            existing_indexes = self.pc.list_indexes()
            index_names = [idx['name'] for idx in existing_indexes]

            if settings.PINECONE_INDEX_NAME not in index_names:
                print(f"Creating Pinecone index: {settings.PINECONE_INDEX_NAME}")
                self.pc.create_index(
                    name=settings.PINECONE_INDEX_NAME,
                    dimension=settings.PINECONE_DIMENSION,
                    metric='cosine',
                    spec=ServerlessSpec(
                        cloud='aws',
                        region='us-east-1'
                    )
                )
                print("✅ Pinecone index created")
            else:
                print(f"✅ Using existing Pinecone index: {settings.PINECONE_INDEX_NAME}")

            # Initialize vector store — this is our single source of truth for retrieval
            self.vector_store = PineconeVectorStore(
                index_name=settings.PINECONE_INDEX_NAME,
                embedding=self.embeddings
            )

            # Expose vector_store directly as retriever
            # RAGEngine.similarity_search_with_score() works on PineconeVectorStore
            self.retriever = self.vector_store

            print("✅ Pinecone vector store ready — using direct dense retrieval")

        except Exception as e:
            print(f"⚠️ Pinecone initialization error: {e}")
            self.vector_store = None
            self.retriever = None

    def _load_metadata(self):
        """Load document metadata from JSON file"""
        metadata_path = settings.BASE_DIR / "metadata.json"
        if metadata_path.exists():
            try:
                with open(metadata_path, 'r') as f:
                    metadata_dict = json.load(f)
                    for doc_id, doc_data in metadata_dict.items():
                        doc_data['upload_date'] = datetime.fromisoformat(doc_data['upload_date'])
                        self.document_metadata[doc_id] = DocumentInfo(**doc_data)
                print(f"✅ Loaded metadata for {len(self.document_metadata)} documents")
            except Exception as e:
                print(f"⚠️ Could not load metadata: {e}")

    def _save_metadata(self):
        """Save document metadata to JSON file"""
        metadata_path = settings.BASE_DIR / "metadata.json"
        try:
            metadata_dict = {}
            for doc_id, doc_info in self.document_metadata.items():
                doc_dict = doc_info.model_dump()
                doc_dict['upload_date'] = doc_info.upload_date.isoformat()
                metadata_dict[doc_id] = doc_dict

            with open(metadata_path, 'w') as f:
                json.dump(metadata_dict, f, indent=2)
            print(f"✅ Metadata saved for {len(self.document_metadata)} documents")
        except Exception as e:
            print(f"⚠️ Could not save metadata: {e}")

    # ── Throttled Embedding ──────────────────────────────────────────────

    def add_documents_throttled(self, docs, batch_size: int = 64, max_retries: int = 8):
        """
        Add documents to Pinecone in batches, backing off on rate-limit (429) errors.

        Pinecone's hosted embedding model has a per-minute token budget; on the free
        plan a large upload can exceed it. Instead of failing the whole upload, we
        embed in batches and, on a 429, wait for the budget to refill and retry.
        """
        import time
        total = len(docs)
        for start in range(0, total, batch_size):
            batch = docs[start:start + batch_size]
            for attempt in range(max_retries):
                try:
                    self.vector_store.add_documents(batch)
                    break
                except Exception as e:
                    msg = str(e)
                    rate_limited = any(s in msg for s in ("429", "RESOURCE_EXHAUSTED", "Too Many Requests"))
                    if rate_limited and attempt < max_retries - 1:
                        wait = 60  # per-minute budget — wait a full window to refill
                        print(f"⏳ Embedding rate limit hit — waiting {wait}s, then resuming "
                              f"(chunks {start + 1}-{min(start + batch_size, total)} of {total})")
                        time.sleep(wait)
                    else:
                        raise
            print(f"  ✅ Embedded {min(start + batch_size, total)}/{total} chunks")

    # ── Text Extraction ─────────────────────────────────────────────────

    def extract_text(self, file_path: Path, file_type: DocumentType) -> str:
        """Extract text from various document formats"""
        try:
            if file_type == DocumentType.PDF:
                return self._extract_pdf(file_path)
            elif file_type == DocumentType.TXT:
                return self._extract_text(file_path)
            elif file_type == DocumentType.DOCX:
                return self._extract_docx(file_path)
            elif file_type == DocumentType.CSV:
                return self._extract_csv(file_path)
            elif file_type == DocumentType.XLSX:
                return self._extract_excel(file_path)
            elif file_type == DocumentType.MARKDOWN:
                return self._extract_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF with page-level structure preservation"""
        text = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text()
                if page_text.strip():
                    text.append(f"[Page {page_num}]\n{page_text}")
        return "\n\n".join(text)

    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX with structure"""
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    parts.append(f"\n## {para.text}\n")
                else:
                    parts.append(para.text)
        return "\n\n".join(parts)

    def _extract_csv(self, file_path: Path) -> str:
        """Extract and format CSV data with semantic structure"""
        df = pd.read_csv(file_path)
        text_parts = [f"CSV Data from {file_path.name}:"]
        text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
        text_parts.append(f"Total rows: {len(df)}\n")

        for idx, row in df.iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            text_parts.append(f"Row {idx + 1}: {row_text}")

        return "\n".join(text_parts)

    def _extract_excel(self, file_path: Path) -> str:
        """Extract and format Excel data with sheet-level structure"""
        xls = pd.ExcelFile(file_path)
        text_parts = [f"Excel Data from {file_path.name}:"]

        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            text_parts.append(f"\n--- Sheet: {sheet_name} ---")
            text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
            text_parts.append(f"Total rows: {len(df)}\n")

            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(f"Row {idx + 1}: {row_text}")

        return "\n".join(text_parts)

    # ── Text Cleaning ────────────────────────────────────────────────────

    def clean_text(self, text: str) -> str:
        """Precision cleaning — preserve meaningful content, remove noise"""
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        text = re.sub(r'\n\s*\d{1,3}\s*\n', '\n', text)
        text = re.sub(r'[\.]{4,}', '...', text)
        text = re.sub(r'[-]{4,}', '---', text)
        return text.strip()

    # ── Document Processing ──────────────────────────────────────────────

    async def process_document(self, file_path: Path, filename: str) -> DocumentInfo:
        """Process a document: extract, chunk, embed, and store in Pinecone"""
        doc_id = str(uuid.uuid4())

        try:
            # Determine file type
            file_ext = file_path.suffix.lower().lstrip('.')
            file_type = DocumentType(file_ext)

            # Create document info
            doc_info = DocumentInfo(
                id=doc_id,
                filename=filename,
                file_type=file_type,
                file_size=file_path.stat().st_size,
                upload_date=datetime.now(),
                status=ProcessingStatus.PROCESSING
            )

            # Extract text
            text = self.extract_text(file_path, file_type)

            if not text.strip():
                raise Exception("No text could be extracted from document")

            cleaned_text = self.clean_text(text)

            # Metadata enrichment
            category = "general"
            text_lower = cleaned_text[:2000].lower()
            if "policy" in text_lower or "guideline" in text_lower:
                category = "policy"
            elif "manual" in text_lower or "instruction" in text_lower:
                category = "manual"
            elif "report" in text_lower:
                category = "report"

            base_metadata = {
                "document_id": doc_id,
                "filename": filename,
                "file_type": file_type.value,
                "source": "document",         # IMPORTANT: mark as document (not chat)
                "category": category,
                "ingested_at": datetime.now().isoformat()
            }

            # Heading-aware chunking
            if "##" in cleaned_text or any(line.isupper() for line in cleaned_text.split("\n")):
                print("📑 Using heading-aware chunking")
                sections = self.split_by_headings(cleaned_text)
            else:
                sections = [{"section_title": "general", "content": cleaned_text}]

            chunker = SmartSemanticChunker(
                min_tokens=100,
                max_tokens=500,
                target_tokens=400,
                overlap_tokens=75
            )

            raw_chunks = []

            for section in sections:
                section_chunks = chunker.chunk_text(section["content"], base_metadata)
                for ch in section_chunks:
                    ch["metadata"]["section_title"] = section["section_title"]
                    ch["metadata"]["section"] = section["section_title"]
                raw_chunks.extend(section_chunks)

            lc_docs = []
            for rc in raw_chunks:
                rc_meta = rc["metadata"]
                rc_meta["doc_id"] = doc_id
                if "section_title" not in rc_meta:
                    rc_meta["section_title"] = category

                # Enrich text with title + section + keywords so that embedding
                # captures document context for better retrieval match
                title_prefix = f"Title: {filename} | Section: {rc_meta.get('section_title', '')}"
                keywords_prefix = f"Keywords: {', '.join(rc_meta.get('keywords', []))}"
                enriched_text = f"{title_prefix}\n{keywords_prefix}\n\n{rc['text']}"

                lc_docs.append(LangChainDocument(page_content=enriched_text, metadata=rc_meta))

            if self.vector_store and lc_docs:
                self.add_documents_throttled(lc_docs)
                print(f"✅ Indexed {len(lc_docs)} chunks into Pinecone for '{filename}'")

            doc_info.chunk_count = len(lc_docs)
            doc_info.status = ProcessingStatus.COMPLETED

            # Save metadata
            self.document_metadata[doc_id] = doc_info
            self._save_metadata()

            return doc_info

        except Exception as e:
            doc_info.status = ProcessingStatus.FAILED
            doc_info.error_message = str(e)
            self.document_metadata[doc_id] = doc_info
            self._save_metadata()
            raise

    async def delete_document(self, doc_id: str):
        """Delete a document from Pinecone and metadata"""
        if doc_id not in self.document_metadata:
            raise Exception(f"Document {doc_id} not found")

        doc_info = self.document_metadata[doc_id]
        file_ext = Path(doc_info.filename).suffix
        file_path = settings.UPLOAD_DIR / f"{doc_id}{file_ext}"

        # 1. Delete from Pinecone
        if self.vector_store:
            try:
                index = self.pc.Index(settings.PINECONE_INDEX_NAME)
                index.delete(filter={"document_id": doc_id})
                print(f"✅ Deleted vectors for {doc_id} from Pinecone")
            except Exception as e:
                print(f"⚠️ Failed to delete from Pinecone: {e}")

        # 2. Delete local file
        if file_path.exists():
            file_path.unlink()
            print(f"✅ Deleted local file {file_path}")

        # 3. Update metadata
        del self.document_metadata[doc_id]
        self._save_metadata()

        return True

    def get_document_count(self) -> int:
        """Get total number of processed documents"""
        return len(self.document_metadata)

    def list_documents(self) -> List[DocumentInfo]:
        """List all processed documents"""
        return list(self.document_metadata.values())

    def get_vector_store(self):
        """Get the Pinecone vector store instance"""
        return self.vector_store

    def get_retriever(self):
        """Get the retriever (same as vector_store for direct dense retrieval)"""
        return self.vector_store


# Global document processor instance
document_processor = DocumentProcessor()
